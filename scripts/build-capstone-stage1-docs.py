from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
CAPSTONE = ROOT / "CAPSTONE"
SOURCE = CAPSTONE / "source" / "bus311-capstone.json"

NAVY = "152536"
TERRA = "9C4A2B"
GOLD = "C18A42"
STEEL = "355773"
SAGE = "4A7C5E"
CREAM = "FAF8F3"
PALE_TERRA = "F7EEE9"
PALE_GOLD = "FFF4DF"
PALE_BLUE = "EDF3F7"
BORDER = "D8D3C8"
GRAY = "5B6472"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, *, size=10.5, color="1A1F2C", bold=False, italic=False, name="Aptos") -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=8.5, color=GRAY)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def configure_document(doc: Document, *, title: str, source_version: str, source_hash: str) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 13, 6),
        ("Heading 2", 13, STEEL, 10, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.2)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    clear_paragraph(header)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("BUS311 Corporate Finance  |  Individual Company Capstone")
    set_run_font(run, size=8.5, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    clear_paragraph(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)

    doc.core_properties.title = title
    doc.core_properties.subject = "BUS311 Capstone Stage 1"
    doc.core_properties.author = "Professor Bethany Evitts"
    doc.core_properties.comments = (
        f"Derived from CAPSTONE/source/bus311-capstone.json v{source_version}; "
        f"SHA-256 {source_hash}"
    )
    doc.core_properties.version = source_version


def add_masthead(doc: Document, *, kicker: str, title: str, source_version: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=9.5, color=TERRA, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=22, color=NAVY, bold=True, name="Aptos Display")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"5 points  |  Fall 2026  |  Approved student release  |  Source v{source_version}")
    set_run_font(run, size=9.25, color=GRAY)

    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GOLD)
    border.append(bottom)
    p_pr.append(border)
    p.paragraph_format.space_after = Pt(8)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, italic=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_lead):])
        set_run_font(run, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)


def add_bullet(doc: Document, text: str, *, checkbox=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    run = p.add_run(("☐ " if checkbox else "• ") + text)
    set_run_font(run)


def remove_empty_page_break_before(doc: Document, heading: str) -> None:
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != heading or index == 0:
            continue
        previous = paragraphs[index - 1]
        breaks = previous._p.findall(
            ".//w:br",
            {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if not previous.text.strip() and breaks:
            previous._p.getparent().remove(previous._p)
        return
    raise ValueError(f"Heading not found while checking page break: {heading}")


def add_step(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)


def add_callout(doc: Document, label: str, text: str, *, fill=PALE_GOLD, accent=TERRA) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label.upper()}  ")
    set_run_font(run, size=9.5, color=accent, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=NAVY)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, rows: list[list[str]], widths: list[int], *, font_size=9.2) -> None:
    table = doc.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.03
            run = p.add_run(value)
            set_run_font(
                run,
                size=font_size,
                color=NAVY if row_index == 0 else "1A1F2C",
                bold=row_index == 0,
            )
            if row_index == 0:
                set_cell_shading(cell, PALE_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, CREAM)
    mark_header_row(table.rows[0])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_response_line(doc: Document, label: str, prompt="[Type here.]") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(f"{label}: ")
    set_run_font(run, bold=True, color=NAVY)
    run = p.add_run(prompt)
    set_run_font(run, color=GRAY, italic=True)


def build_stage_one_menu(capstone: dict, source_hash: str) -> Path:
    stage = next(item for item in capstone["hub"]["stages"] if item["stageId"] == "S01_SCOPE")
    output = CAPSTONE / "bus311-capstone-cfo-decision-menu.docx"
    doc = Document()
    configure_document(
        doc,
        title="BUS311 Stage 1 - Company Research and Potential CFO Decision Starter",
        source_version=capstone["meta"]["sourceVersion"],
        source_hash=source_hash,
    )
    add_masthead(
        doc,
        kicker="BUS311 Capstone Stage 1",
        title="Company Research and Potential CFO Decision Starter",
        source_version=capstone["meta"]["sourceVersion"],
    )

    add_body(doc, stage["objective"])
    add_callout(
        doc,
        "Stage 1 is a guided first look",
        "You are choosing a company, noticing what may matter, and proposing one possible direction. FactSet is required. You do not need to prove the idea or make a final recommendation yet.",
    )

    doc.add_paragraph("What you submit", style="Heading 1")
    for item in (
        "The approved public company's name, ticker, and U.S. exchange.",
        "A short beginner-friendly company snapshot based on AI and/or ordinary web search.",
        "One concrete thing learned from FactSet, plus the screen, report, or feature used.",
        "One potential company-specific CFO decision and a short reason it seems worth investigating.",
        "Two questions you want to research in Stage 2.",
    ):
        add_bullet(doc, item)

    doc.add_paragraph("A five-step start", style="Heading 1")
    for item in (
        "Choose one approved public company and record its basic identifying information.",
        "Use ordinary web search, AI, or both to learn what the company sells, who it serves, and what issues or opportunities appear important.",
        "Open FactSet and explore at least one company screen, report, or feature.",
        "Write one specific FactSet learning in your own words and name where you found it.",
        "Choose one potential CFO decision from the menu below. Treat it as a starting idea that may change.",
    ):
        add_step(doc, item)

    doc.add_paragraph("FactSet is required", style="Heading 1")
    add_body(doc, stage["factSetLearning"])
    add_table(
        doc,
        [
            ["A beginner-appropriate FactSet discovery", "Example of what to record"],
            ["Company or segment metric", "A metric or segment result that helped you understand the business."],
            ["Ownership or market item", "An ownership, trading, capital-markets, or security detail that caught your attention."],
            ["Peer or comparison observation", "One way the company looked similar to or different from a relevant peer."],
            ["Estimate", "An analyst estimate or expectation that raised a useful question."],
            ["News or data discovery", "A relevant FactSet news item, event, chart, or data feature you did not know about before."],
        ],
        [2700, 6660],
    )
    add_callout(
        doc,
        "Keep it simple",
        "For Stage 1, name the FactSet screen, report, or feature and explain what you learned. You do not need a formal citation, definition audit, reconciliation, or independent validation yet.",
        fill=PALE_BLUE,
        accent=STEEL,
    )

    doc.add_paragraph("Potential CFO decision directions", style="Heading 1")
    add_body(
        doc,
        "Browse the ten directions and choose one that seems connected to your initial research. You are not comparing or defending alternatives in Stage 1.",
    )
    rows = [["Decision area", "Beginner starting question"]]
    rows.extend([[item["area"], item["prompt"]] for item in capstone["hub"]["decisionMenu"]])
    add_table(doc, rows, [2200, 7160], font_size=8.8)

    doc.add_paragraph("Write the potential decision in your own words", style="Heading 1")
    add_callout(
        doc,
        "Sentence starter",
        "One potential decision the CFO might need to consider is whether [company] should [possible action], because my initial research made me curious about [company-specific issue or opportunity].",
        fill=PALE_TERRA,
    )

    doc.add_paragraph("What comes later", style="Heading 1")
    add_table(
        doc,
        [
            ["Stage 1", "Later stages"],
            ["Explore the company and notice a potential CFO decision.", "Build a traceable evidence register and analyze the revenue engine."],
            ["Record one concrete FactSet learning and the feature used.", "Add full FactSet labels, dates, units, definitions, and limitations."],
            ["Keep the decision tentative.", "Test a hypothesis, compare alternatives, model scenarios, and defend the final recommendation."],
        ],
        [3300, 6060],
    )

    doc.add_paragraph("Before you submit", style="Heading 1")
    for item in (
        "Company name, ticker, and exchange are complete.",
        "Introductory company notes show AI and/or ordinary web research.",
        "The required FactSet learning is specific and names the screen, report, or feature used.",
        "One potential CFO decision is written in your own words.",
        "The brief treats the idea as exploratory rather than proven or final.",
    ):
        add_bullet(doc, item, checkbox=True)

    doc.save(output)
    return output


def build_stage_one_brief(capstone: dict, source_hash: str) -> Path:
    stage = next(item for item in capstone["hub"]["stages"] if item["stageId"] == "S01_SCOPE")
    output = CAPSTONE / "bus311-capstone-stage-1-exploration-brief.docx"
    doc = Document()
    configure_document(
        doc,
        title="BUS311 Stage 1 - Company Exploration Brief",
        source_version=capstone["meta"]["sourceVersion"],
        source_hash=source_hash,
    )
    add_masthead(
        doc,
        kicker="BUS311 Capstone Stage 1",
        title="Company Exploration Brief",
        source_version=capstone["meta"]["sourceVersion"],
    )
    add_body(doc, "Complete this brief in your own words. Keep answers concise and beginner-friendly.")
    add_callout(
        doc,
        "Exploratory, not final",
        "Choose an approved public company, learn the basics through AI and/or ordinary web search, use FactSet, and propose one potential CFO decision. Your idea may change as you learn more.",
    )

    doc.add_paragraph("A. Student and company", style="Heading 1")
    add_table(
        doc,
        [
            ["Field", "Student entry"],
            ["Student name", "[Type here.]"],
            ["Course section", "[Type here.]"],
            ["Date", "[Type here.]"],
            ["Company legal name", "[Type here.]"],
            ["Ticker and U.S. exchange", "[Type here.]"],
            ["Industry or sector", "[Type here.]"],
        ],
        [2700, 6660],
    )

    doc.add_paragraph("B. Introductory company research", style="Heading 1")
    add_body(doc, "Research path used:  ☐ AI  ☐ Ordinary web search  ☐ Both")
    add_body(
        doc,
        "In three to five short bullets, explain what the company sells, who its customers are, the segments or markets that seem important, and one current issue or opportunity you noticed.",
    )
    for _ in range(4):
        add_bullet(doc, "[Type one observation here.]")
    add_response_line(doc, "Websites, AI tool, or other places explored (names only; formal citations are not required yet)")
    add_callout(
        doc,
        "AI privacy reminder",
        "If you use AI, provide only a short student-written summary. Do not upload FactSet files or screenshots, a workbook, personal information, instructor-only material, or nonpublic information.",
        fill=PALE_BLUE,
        accent=STEEL,
    )

    doc.add_paragraph("C. Required FactSet discovery", style="Heading 1")
    add_callout(
        doc,
        "FactSet is required",
        "Explore one company screen, report, or feature. Record one specific thing you learned. A formal citation or validation is not required in Stage 1.",
        fill=PALE_GOLD,
        accent=TERRA,
    )
    add_table(
        doc,
        [
            ["Required field", "Student entry"],
            ["FactSet screen, report, or feature used", "[Type the name shown in FactSet.]"],
            ["One specific thing I learned from FactSet", "[State the metric, observation, estimate, comparison, ownership/market item, or news/data discovery in your own words.]"],
            ["Why it caught my attention", "[Explain briefly how it helped you understand the company or raised a useful question.]"],
        ],
        [3100, 6260],
    )
    add_body(
        doc,
        "Examples that work: a company or segment metric, an ownership or market item, a peer/comparison observation, an estimate, or a relevant news/data discovery.",
        italic=True,
    )

    doc.add_paragraph("D. One potential CFO decision", style="Heading 1")
    add_response_line(
        doc,
        "Decision area",
        "[Choose one: growth; capital allocation; financing; payout; acquisition/partnership/divestiture; operations/working capital; risk; other.]",
    )
    add_response_line(
        doc,
        "Potential decision",
        "One potential decision the CFO might need to consider is whether [company] should [possible action].",
    )
    add_response_line(
        doc,
        "Why this seems worth investigating",
        "[In 50-100 words, connect the possible decision to something you noticed in your introductory research. You are explaining curiosity, not proving a claim.]",
    )
    add_response_line(doc, "Question I want to investigate in Stage 2", "[Type here.]")
    add_response_line(doc, "A second question that could help me learn more", "[Type here.]")

    doc.add_paragraph("E. Stage 1 submission check", style="Heading 1")
    for item in (
        "I selected an approved public company and entered its ticker and exchange.",
        "I used AI and/or ordinary web search for introductory company research.",
        "I used FactSet, stated one concrete learning, and named the screen, report, or feature used.",
        "I proposed one potential CFO decision in my own words.",
        "I understand that this is a tentative starting point that may change.",
    ):
        add_bullet(doc, item, checkbox=True)
    add_callout(
        doc,
        "Not required in Stage 1",
        "Do not verify the recommendation, validate a hypothesis, compare or defend alternatives, build an evidence register, prove the claim, complete formal citations, or seek a formal approval gate yet. Those more rigorous activities belong in later stages.",
        fill=PALE_TERRA,
        accent=TERRA,
    )

    doc.save(output)
    return output


def iter_paragraphs(doc: Document):
    seen_cells: set[object] = set()
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                marker = cell._tc
                if marker in seen_cells:
                    continue
                seen_cells.add(marker)
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                yield paragraph
    for section in doc.sections:
        for part in (section.header, section.footer, section.even_page_header, section.even_page_footer):
            for paragraph in part.paragraphs:
                yield paragraph


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_text(doc: Document, old: str, new: str, *, required=True) -> int:
    count = 0
    for paragraph in iter_paragraphs(doc):
        if old in paragraph.text:
            set_paragraph_text(paragraph, paragraph.text.replace(old, new))
            count += 1
    if required and count == 0:
        raise ValueError(f"required DOCX text not found: {old}")
    return count


def set_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_field(paragraph)


# The full assignment, student rubric, AI guide, and prompt guide are maintained by
# scripts/build-capstone-student-docs.py. This builder now owns only the two Stage 1 files.


def main() -> int:
    source_bytes = SOURCE.read_bytes()
    capstone = json.loads(source_bytes)
    source_hash = hashlib.sha256(source_bytes).hexdigest()
    outputs = [
        build_stage_one_menu(capstone, source_hash),
        build_stage_one_brief(capstone, source_hash),
    ]
    for output in outputs:
        print(output.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
