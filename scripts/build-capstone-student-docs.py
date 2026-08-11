from __future__ import annotations

import hashlib
import importlib.util
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent.parent
CAPSTONE = ROOT / "CAPSTONE"
SOURCE = CAPSTONE / "source" / "bus311-capstone.json"

helper_spec = importlib.util.spec_from_file_location(
    "capstone_doc_helpers", ROOT / "scripts" / "build-capstone-stage1-docs.py"
)
helpers = importlib.util.module_from_spec(helper_spec)
assert helper_spec.loader is not None
helper_spec.loader.exec_module(helpers)


def add_title(doc: Document, kicker: str, title: str, meta: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    helpers.set_run_font(p.add_run(kicker.upper()), size=9.5, color=helpers.TERRA, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    helpers.set_run_font(p.add_run(title), size=24, color=helpers.NAVY, bold=True, name="Aptos Display")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    helpers.set_run_font(p.add_run(meta), size=9.25, color=helpers.GRAY)
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), helpers.GOLD)
    border.append(bottom)
    p_pr.append(border)
    p.paragraph_format.space_after = Pt(8)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_bullets(doc: Document, items: list[str], *, checkbox: bool = False) -> None:
    for item in items:
        helpers.add_bullet(doc, item, checkbox=checkbox)


def set_properties(doc: Document, capstone: dict, source_hash: str, *, title: str, rubric: bool = False) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = "BUS311 Individual Company Capstone"
    doc.core_properties.author = "Professor Bethany Evitts"
    doc.core_properties.comments = (
        f"Derived from CAPSTONE/source/bus311-capstone.json v{capstone['meta']['sourceVersion']}; "
        f"SHA-256 {source_hash}"
    )
    doc.core_properties.version = capstone["meta"]["sourceVersion"]
    helpers.set_approved_footer(doc, capstone["meta"]["sourceVersion"], source_hash, rubric=rubric)


def build_assignment(capstone: dict, source_hash: str) -> Path:
    doc = Document()
    helpers.configure_document(
        doc,
        title="BUS311 Individual Company Capstone Assignment",
        source_version=capstone["meta"]["sourceVersion"],
        source_hash=source_hash,
    )
    add_title(
        doc,
        "BUS311 Corporate Finance · Approved student assignment",
        capstone["project"]["title"],
        f"100 points  |  25% of BUS311  |  Individual project  |  Fall 2026  |  Source v{capstone['meta']['sourceVersion']}",
    )
    helpers.add_callout(doc, "Board assignment", capstone["hub"]["purpose"], fill=helpers.PALE_GOLD, accent=helpers.TERRA)

    add_heading(doc, "At a glance", 1)
    helpers.add_table(
        doc,
        [
            ["Item", "Approved requirement"],
            ["Audience", capstone["project"]["audience"]],
            ["Company", capstone["project"]["companyEligibility"]],
            ["Decision", capstone["project"]["decision"]],
            ["Stages 1-4", "5 + 8 + 8 + 4 = 25 points; dates and requirements are unchanged"],
            ["Stage 5 component 1", "PowerPoint company-analysis project submission — 50 points"],
            ["Stage 5 component 2", "oral Board presentation to the class — 25 points"],
            ["Stage 5 upload", f"Exactly one editable .pptx due {capstone['finalSubmission']['deadlineLabel']}"],
            ["Live presentation", "On the assigned class date; up to 7 minutes presenting plus up to 3 minutes responding to live questions, scored together"],
            ["Project total", "100 points = 25% of the BUS311 course grade"],
        ],
        [2500, 6860],
        font_size=9.25,
    )

    add_heading(doc, "The decision you will build", 1)
    helpers.add_body(doc, capstone["hub"]["centralQuestion"])
    add_bullets(doc, [
        "Lead with one material, company-specific action the CFO could recommend to the Board.",
        "Diagnose the revenue engine before selecting forecast and valuation assumptions.",
        "Use traceable SEC, permitted FactSet, Excel, and other credible evidence beginning in Stage 2.",
        "Compare credible alternatives and specify two or three supporting actions, owners, timing, resources, metrics, triggers, and risks.",
        "Complete the work individually and be prepared to explain every submitted figure, source choice, formula, assumption, and recommendation.",
    ])

    add_heading(doc, "Semester schedule and points", 1)
    milestone_rows = [["Stage", "Deliverable", "Due", "Points"]]
    for milestone, stage in zip(capstone["milestones"], capstone["hub"]["stages"][:4]):
        milestone_rows.append([stage["label"], milestone["title"], milestone["dueLabel"], str(milestone["points"])])
    milestone_rows.extend([
        ["Stage 5 upload", "PowerPoint company-analysis project submission", capstone["finalSubmission"]["deadlineLabel"], "50"],
        ["Stage 5 live", "oral Board presentation to the class", "Assigned date: Nov. 30, Dec. 2, Dec. 7, or Dec. 9", "25"],
        ["", "Project total", "", "100"],
    ])
    helpers.add_table(doc, milestone_rows, [1250, 3330, 3480, 1300], font_size=8.7)

    add_heading(doc, "Stage 1 — Explore the company and a CFO decision (5 points)", 1)
    stage = capstone["hub"]["stages"][0]
    helpers.add_callout(doc, "Outcome", stage["outcome"], fill=helpers.PALE_BLUE, accent=helpers.STEEL)
    helpers.add_body(doc, f"Due: {capstone['milestones'][0]['dueLabel']}", bold_lead="Due: ")
    helpers.add_body(doc, "Submit: Stage 1 company exploration brief.", bold_lead="Submit: ")
    add_bullets(doc, [
        "Identify the approved company, ticker, exchange, and your individual ownership of the work.",
        "Use AI and/or ordinary web search for a beginner-friendly company snapshot; treat these notes as exploratory rather than verified evidence.",
        stage["factSetLearning"],
        "Propose one potential company-specific CFO decision in your own words and explain why it seems worth investigating.",
        "List two questions to investigate in Stage 2.",
    ])
    helpers.add_callout(doc, "Not required yet", "; ".join(stage["notRequiredYet"]) + ".", fill=helpers.PALE_GOLD, accent=helpers.TERRA)

    add_heading(doc, "Stage 2 — Explain the revenue engine (8 points)", 1)
    stage = capstone["hub"]["stages"][1]
    helpers.add_callout(doc, "Outcome", stage["outcome"], fill=helpers.PALE_BLUE, accent=helpers.STEEL)
    helpers.add_body(doc, f"Due: {capstone['milestones'][1]['dueLabel']}", bold_lead="Due: ")
    helpers.add_body(doc, f"Submit: {stage['requiredSubmission']}.", bold_lead="Submit: ")
    add_bullets(doc, [
        "Map relevant products or services, segments, geographies, customers, channels, pricing, and recurring or transactional revenue.",
        "Analyze three to five comparable fiscal years and reconcile changes in periods, units, definitions, and segment reporting.",
        "Separate company-relevant drivers such as price, volume, mix, currency, acquisitions, capacity, retention, utilization, or backlog from temporary effects.",
        "Use SEC and permitted FactSet evidence with source, period, units, definition, exact location, and limitations.",
        "State a testable revenue hypothesis naming the driver, expected effect, period, and falsification evidence.",
        "Verify the strongest retained AI challenge and record an accept, modify, or reject judgment.",
    ])

    add_heading(doc, "Stage 3 — Model value and scenarios (8 points)", 1)
    stage = capstone["hub"]["stages"][2]
    helpers.add_callout(doc, "Outcome", stage["outcome"], fill=helpers.PALE_BLUE, accent=helpers.STEEL)
    helpers.add_body(doc, f"Due: {capstone['milestones'][2]['dueLabel']}", bold_lead="Due: ")
    helpers.add_body(doc, f"Submit: {stage['requiredSubmission']}.", bold_lead="Submit: ")
    add_bullets(doc, [
        "Separate inputs, sources, calculations, outputs, checks, scenarios, and sensitivities.",
        "Link base, upside, and downside cases explicitly to revenue drivers and the developing CFO decision.",
        "Use applicable tax, NOPAT, depreciation, capital expenditure, working-capital, discount-rate, terminal-value, and enterprise-to-equity logic.",
        "Resolve formula and sign errors; show scenario consistency, valuation bridges, terminal-value reasonableness, and sensitivity links.",
        "Label every material SEC and FactSet assumption with the item or report name, as-of date, units, and limitations.",
    ])

    add_heading(doc, "Stage 4 — Challenge and revise (4 points)", 1)
    stage = capstone["hub"]["stages"][3]
    helpers.add_callout(doc, "Outcome", stage["outcome"], fill=helpers.PALE_BLUE, accent=helpers.STEEL)
    helpers.add_body(doc, f"Due: {capstone['milestones'][3]['dueLabel']}", bold_lead="Due: ")
    helpers.add_body(doc, f"Submit: {stage['requiredSubmission']}.", bold_lead="Submit: ")
    add_bullets(doc, [
        "Document context, assumptions, prompt purpose, strongest challenge, and the output retained.",
        "Check the challenge using an exact SEC filing, FactSet item, Excel calculation, or course concept.",
        "Record the finding, uncertainty, accept/modify/reject judgment, and the resulting revision or defense.",
        "Prepare the final verified AI-use disclosure and rehearse the oral Board presentation, including live questions.",
        "Never provide raw licensed FactSet files or screenshots, private information, instructor-only material, or nonpublic information to a public AI tool.",
    ])

    doc.add_page_break()
    add_heading(doc, "Stage 5 — Submit the company analysis and present to the Board (75 points)", 1)
    helpers.add_callout(doc, "Exactly two assessed components", "PowerPoint company-analysis project submission — 50 points; oral Board presentation to the class — 25 points. There is one uploaded file and one live performance.", fill=helpers.PALE_GOLD, accent=helpers.TERRA)
    helpers.add_table(
        doc,
        [
            ["Expectation", "PowerPoint upload", "Live oral presentation"],
            ["What you do", "Upload one editable company-analysis .pptx", "Present the submitted deck individually to the class"],
            ["When", capstone["finalSubmission"]["deadlineLabel"], "Assigned date between Nov. 30 and Dec. 9"],
            ["Length", "8-10 core slides plus optional analytical appendix", "Up to 7 minutes presenting plus up to 3 minutes responding to live questions"],
            ["Points", "50", "25"],
            ["Canvas", "One file is uploaded", "No second Stage 5 file is uploaded"],
        ],
        [1900, 3730, 3730],
        font_size=8.75,
    )

    add_heading(doc, "Component 1 — PowerPoint company-analysis project submission (50 points)", 2)
    helpers.add_body(doc, f"Required filename: {capstone['finalSubmission']['files'][0]}", bold_lead="Required filename: ")
    add_bullets(doc, [
        "Submit exactly one editable PowerPoint (.pptx) by the final file lock. PDF-only, image-only, link-only, or noneditable substitutes do not meet the format requirement.",
        "Use 8-10 core slides. Optional analytical appendix slides may support deeper evidence and live questions; they do not replace a concise core decision narrative.",
        "Lead with the requested Board decision and one company-specific CFO recommendation.",
        "Show the revenue engine, three-to-five-year evidence, company-specific drivers, testable hypothesis, and verified red-team revision before valuation.",
        "Reconcile valuation or capital-budgeting outputs, base/upside/downside scenarios, and decision-changing sensitivities to the assessed Excel model.",
        "Compare credible alternatives and specify two or three actions with owners, timing, resources, dependencies, indicators, triggers, risks, and mitigations.",
        "Use readable, accessible, editable or traceable visuals with units, periods, definitions, compact SEC/FactSet/model source notes, limitations, and the verified AI-use disclosure.",
    ])

    doc.add_page_break()
    add_heading(doc, "Component 2 — oral Board presentation to the class (25 points)", 2)
    helpers.add_callout(doc, "Live requirement", "Present the submitted PowerPoint individually on your assigned date. The live performance is not a second upload. Delivery and responses to questions are assessed together in this single 25-point criterion.", fill=helpers.PALE_BLUE, accent=helpers.STEEL)
    add_bullets(doc, [
        "Use up to seven minutes for the decision briefing. Lead with the recommendation, explain revenue before valuation, and use the deck as evidence rather than as a script.",
        "Use a clear professional voice, controlled pacing, purposeful transitions, accurate terminology, and communication appropriate for a CFO-and-Board audience in the classroom.",
        "For up to three additional minutes, respond to live questions about sources, revenue drivers, assumptions, calculations, scenarios, alternatives, implementation, risks, and conditions that would change the decision.",
        "Trace answers to SEC, FactSet, Excel, or course evidence. Reconcile differing definitions when challenged.",
        "If uncertain, state the limitation and identify the next evidence or calculation needed rather than inventing an answer.",
        "Demonstrate individual ownership of the research, model, recommendation, submitted deck, and live responses.",
    ])
    helpers.add_table(
        doc,
        [
            ["Class date", "Planned presenters", "Individual time target"],
            ["Nov. 30", "5", "Up to 7 minutes presenting + up to 3 minutes for live questions + transition"],
            ["Dec. 2", "5", "Same"],
            ["Dec. 7", "5", "Same"],
            ["Dec. 9", "4", "Same"],
        ],
        [2100, 2100, 5160],
        font_size=9,
    )

    add_heading(doc, "Scoring overview", 1)
    helpers.add_table(doc, [["Criterion", "Points"]] + [[item["title"], str(item["points"])] for item in capstone["rubric"]["criteria"]] + [["Project total", "100"]], [7960, 1400], font_size=9.1)
    helpers.add_body(doc, "Stages 1-4 remain 5, 8, 8, and 4 points. Stage 5 contains exactly two criteria worth 50 and 25 points.")

    doc.add_page_break()
    add_heading(doc, "Policies and final readiness", 1)
    for label, key in [
        ("Individual work", "individualWork"),
        ("Sources and FactSet", "factSetAndSources"),
        ("Verified AI judgment", "verifiedAIJudgment"),
        ("Late milestones", "lateWork"),
        ("Revision", "revision"),
        ("PowerPoint file lock", "finalFileLock"),
        ("Submission location", "submissionLocation"),
        ("Absence", "absence"),
    ]:
        helpers.add_body(doc, f"{label}: {capstone['policies'][key]}", bold_lead=f"{label}: ")
    add_heading(doc, "Final readiness checklist", 2)
    add_bullets(doc, [
        "My one editable PowerPoint uses the required filename and is uploaded before the file lock.",
        "My 8-10 core slides make the decision, revenue logic, model-linked valuation, alternatives, implementation, risks, limitations, sources, and AI disclosure easy to trace.",
        "I have rehearsed the decision briefing to fit within seven minutes without reading the slides.",
        "I can respond for up to three minutes to live questions about my sources, model, assumptions, scenarios, alternatives, implementation, and risks.",
        "I understand the 25-point oral criterion includes both delivery and responses to live questions.",
        "I know my assigned class presentation date and will use the same PowerPoint submitted by the file lock.",
    ], checkbox=True)

    set_properties(doc, capstone, source_hash, title="BUS311 Individual Company Capstone Assignment")
    output = CAPSTONE / "bus311-capstone-assignment.docx"
    doc.save(output)
    return output


def build_rubric(capstone: dict, source_hash: str) -> Path:
    doc = Document()
    helpers.configure_document(doc, title="BUS311 Individual Company Capstone Student Rubric", source_version=capstone["meta"]["sourceVersion"], source_hash=source_hash)
    add_title(doc, "BUS311 Corporate Finance · Student rubric", "Individual Company Capstone Rubric", f"100 points  |  Complete · Developing · Not demonstrated  |  Source v{capstone['meta']['sourceVersion']}")
    helpers.add_callout(doc, "Point structure", "Stages 1-4 = 25 points (5 + 8 + 8 + 4). Stage 5 has exactly two criteria: PowerPoint company-analysis project submission = 50 points; oral Board presentation to the class = 25 points. Project total = 100 points.", fill=helpers.PALE_GOLD, accent=helpers.TERRA)
    helpers.add_table(doc, [["Component", "Points"], ["Stages 1-4: milestones, Excel model, and revision", "25"], ["PowerPoint company-analysis project submission", "50"], ["oral Board presentation to the class", "25"], ["Project total", "100"]], [7960, 1400], font_size=9.3)
    helpers.add_body(doc, "Use the evidence list under each criterion to prepare. The rating table shows the concrete evidence expected for Complete, Developing, and Not demonstrated performance.")

    for index, criterion in enumerate(capstone["rubric"]["criteria"]):
        doc.add_page_break()
        add_heading(doc, f"{index + 1}. {criterion['title']} — {criterion['points']} points", 1)
        helpers.add_body(doc, criterion["description"])
        add_heading(doc, "Required evidence", 2)
        add_bullets(doc, criterion["evidenceRules"])
        add_heading(doc, "Performance evidence", 2)
        rows = [["Rating", "Points", "Evidence description"]]
        for level_id in ("COMPLETE", "DEVELOPING", "NOT_DEMONSTRATED"):
            level = criterion["performanceLevels"][level_id]
            label = next(item["label"] for item in capstone["rubric"]["performanceScale"] if item["levelId"] == level_id)
            rows.append([label, str(level["points"]), level["description"]])
        helpers.add_table(doc, rows, [1600, 1100, 6660], font_size=8.7)

    doc.add_page_break()
    add_heading(doc, "Final point check", 1)
    helpers.add_table(doc, [["Criterion", "Maximum points"]] + [[item["title"], str(item["points"])] for item in capstone["rubric"]["criteria"]] + [["Total", "100"]], [7960, 1400], font_size=9)
    helpers.add_callout(doc, "Stage 5 scoring boundary", "There is no separate score for delivery or for responses to questions. Both are included in the single 25-point oral Board presentation criterion.", fill=helpers.PALE_BLUE, accent=helpers.STEEL)
    set_properties(doc, capstone, source_hash, title="BUS311 Individual Company Capstone Student Rubric", rubric=True)
    output = CAPSTONE / "bus311-capstone-student-rubric.docx"
    doc.save(output)
    return output


def patch_supporting_docs(capstone: dict, source_hash: str) -> list[Path]:
    changes = {
        "bus311-capstone-ai-student-guide.docx": [
            ("Checkpoint 4 - Board Q&A rehearsal", "Checkpoint 4 - Oral Board presentation rehearsal"),
            ("Revise the deck, executive summary, recommendation logic, or speaking notes as appropriate.", "Revise the PowerPoint company analysis, recommendation logic, model support, or speaking notes as appropriate."),
            ("Board Q&A weakness and the revision made to the recommendation, implementation, risk response, deck, or speaking notes.", "Oral Board presentation weakness and the revision made to the recommendation, implementation, risk response, PowerPoint, or speaking notes."),
            ("The Q&A rehearsal covers revenue, valuation, implementation, and risk.", "The oral-presentation rehearsal covers delivery and live questions about revenue, valuation, implementation, and risk."),
            ("4. Board Q&A", "4. Oral Board presentation"),
            ("Rehearse revenue, valuation, implementation, and risk questions.", "Rehearse the decision briefing and live questions about revenue, valuation, implementation, and risk."),
        ],
        "bus311-capstone-capaj-prompts.docx": [
            ("Checkpoint 4 - Rehearse Board Q&A", "Checkpoint 4 - Rehearse the oral Board presentation"),
            ("Recheck the evidence behind weak answers and identify any revision needed in the deck, executive summary, recommendation, model support, or speaking notes.", "Recheck the evidence behind weak answers and identify any revision needed in the PowerPoint company analysis, recommendation, model support, or speaking notes."),
        ],
    }
    outputs = []
    for name, replacements in changes.items():
        path = CAPSTONE / name
        doc = Document(path)
        for old, new in replacements:
            helpers.replace_text(doc, old, new, required=False)
        all_text = " ".join(p.text for p in helpers.iter_paragraphs(doc)).lower()
        if "executive summary" in all_text or "q&a" in all_text:
            raise ValueError(f"obsolete Stage 5 language remains in {name}")
        set_properties(doc, capstone, source_hash, title=doc.core_properties.title or name)
        doc.save(path)
        outputs.append(path)
    return outputs


def main() -> int:
    source_bytes = SOURCE.read_bytes()
    capstone = json.loads(source_bytes)
    source_hash = hashlib.sha256(source_bytes).hexdigest()
    outputs = [build_assignment(capstone, source_hash), build_rubric(capstone, source_hash)]
    outputs.extend(patch_supporting_docs(capstone, source_hash))
    for output in outputs:
        print(output.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
